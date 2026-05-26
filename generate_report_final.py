from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.1)
    section.right_margin = Inches(1.1)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
style.paragraph_format.space_after = Pt(5)

def eng(text, size=11, bold=False):
    """English term - always bold"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x6E)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    return p

def arb(text, size=11):
    """Arabic explanation only"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    p.paragraph_format.space_after = Pt(6)
    return p

def heading_ar(text, level=1):
    """
    Adds an Arabic heading with the requested level.
    """
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0D, 0x2B, 0x5E)
    return h

def bullet_arb(text, size=11):
    """
    Adds an Arabic bullet paragraph.
    """
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(2)
    return p

def bullet_eng_arb(eng_term, arb_text, size=11):
    """
    Adds a bilingual bullet paragraph.
    """
    p = doc.add_paragraph(style='List Bullet')
    r1 = p.add_run(eng_term)
    r1.bold = True
    r1.font.size = Pt(size)
    r1.font.color.rgb = RGBColor(0x1A, 0x3A, 0x6E)
    r2 = p.add_run(' : ' + arb_text)
    r2.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_table(headers, rows):
    """
    Adds a formatted table to the generated document.
    """
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = cell._element.get_or_add_tcPr()
        shading_elem = shading.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd',
            { '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val': 'clear',
              '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color': 'auto',
              '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill': '1A3A6E' })
        shading.append(shading_elem)
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(9.5)
    doc.add_paragraph()
    return table

# ═══════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
eng('AI Recruiter Assistant', size=28, bold=True)
eng('NLP Pipeline : From CV Upload to Matching + Report', size=14, bold=False)
doc.add_paragraph()
arb('هذا المستند يشرح الـ Pipeline الكامل لتحليل السيرة الذاتية.', 12)
arb('تنقسم المعالجة إلى 6 مراحل متتالية.', 12)
arb('كل مرحلة تستلم Input معين وتُنتج Output للتي تليها.', 12)
doc.add_paragraph()
arb('المصطلحات التقنية باللغة الإنجليزية كما هي.', 11)
arb('الشرح باللغة العربية.', 11)
doc.add_paragraph()
doc.add_paragraph()
eng('- - -', size=14, bold=False)
doc.add_page_break()

# ═══════════════════════════════════════════
# PIPELINE OVERVIEW TABLE
# ═══════════════════════════════════════════
heading_ar('Pipeline Summary', 1)
arb('الجدول التالي يلخص المراحل الست. كل مرحلة تقوم بتحويل معين للبيانات حتى الوصول إلى النتيجة النهائية.')

add_table(
    ['Stage', 'Input', 'Output', 'Key Files'],
    [
        ['1. Text Extraction', 'CV file (PDF/DOCX/TXT)', 'raw_text (plain text)', 'cv_parser.py'],
        ['2. NLP Preprocessing', 'raw_text', 'CandidateProfile (basic)', 'cv_parser.py, stanza_nlp.py, skill_catalog.py'],
        ['3. Enhanced LLM Parsing', 'raw_text + basic profile', 'CandidateProfile (enhanced)', 'enhanced_cv_parser.py, bilingual_llm.py'],
        ['4. Job Parsing', 'Job Description (text)', 'JobProfile', 'job_parser.py, skill_catalog.py'],
        ['5. Matching', 'JobProfile + CandidateProfile[]', 'MatchResult[]', 'hybrid_matcher.py, matching.py, embedding.py'],
        ['6. Report', 'MatchResult', 'CandidateReport', 'explainability.py'],
    ]
)

# ═══════════════════════════════════════════
# STAGE 1: TEXT EXTRACTION
# ═══════════════════════════════════════════
doc.add_page_break()
heading_ar('Stage 1 : Text Extraction', 1)

eng('Objective', bold=True)
arb('الهدف من هذه المرحلة هو تحويل ملف الـ CV من صيغته الأصلية (PDF أو DOCX أو TXT) إلى نص خام يمكن للنظام فهمه.')

eng('Entry Point')
arb('الملف: cv_parser.py')
arb('الدالة: extract_text(file_name, content)')

eng('Supported Formats')
bullet_eng_arb('PDF', 'يستخدم مكتبة pdfplumber لاستخراج النص من الطبقة النصية.')
bullet_eng_arb('PDF (Scanned)', 'إذا كان النص أقل من 10 أحرف، يشغّل OCR باستخدام pytesseract + pdf2image.')
bullet_eng_arb('DOCX', 'يستخدم python-docx لقراءة الفقرات.')
bullet_eng_arb('TXT', 'قراءة مباشرة مع تجاهل أخطاء الترميز.')

eng('Validation')
bullet_arb('يتم التحقق من أن النص الناتج يحتوي على 10 أحرف أبجدية رقمية على الأقل.')
bullet_arb('إذا فشل الاستخراج بالكامل، يُرفع ValueError مع رسالة واضحة.')

eng('Error Handling')
bullet_arb('PDF تالف: "Could not safely extract text from PDF"')
bullet_arb('OCR فاشل: "Could not extract readable text from PDF. OCR returned empty text."')
bullet_arb('صيغة .doc قديمة: "Legacy .doc files are not supported safely"')

eng('Input')
arb('file_name: اسم الملف (jad_cv.pdf)')
arb('content: محتوى الملف بالـ bytes')

eng('Output')
arb('raw_text: نص خام نظيف من الترميزات')
arb('مثال: "Jad Ahmad\\njad.ahmad@email.com\\n+961 3 123456\\nSoftware Engineer with 5 years..."')

# ═══════════════════════════════════════════
# STAGE 2: NLP PREPROCESSING
# ═══════════════════════════════════════════
doc.add_page_break()
heading_ar('Stage 2 : NLP Preprocessing & Parsing', 1)

eng('Objective')
arb('تحويل النص الخام إلى بنية بيانات منظمة تحتوي على الاسم، البريد الإلكتروني، رقم الهاتف، المهارات، الخبرات، والتعليم.')

eng('Entry Point')
arb('الملف: cv_parser.py')
arb('الدالة: parse_cv_text(text)')
arb('الملفات المساعدة: stanza_nlp.py (تقسيم الجمل) و skill_catalog.py (استخراج المهارات)')

eng('2.1 Text Normalization')
arb('الملف: skill_catalog.py')
arb('الدالة: normalize_text_for_skill_matching()')
bullet_arb('تحويل النص إلى lowercase.')
bullet_arb('استبدال / بمسافة.')
bullet_arb('إزالة الرموز غير الأبجدية الرقمية مع الاحتفاظ بـ + # .')
bullet_arb('تطبيع عربي: إ و أ و آ ← ا ، ى ← ي ، ة ← ه ، ؤ ← و ، ئ ← ي.')
bullet_arb('ضغط المسافات المتعددة إلى مسافة واحدة.')

eng('2.2 Sentence Splitting & Tokenization')
arb('الملف: stanza_nlp.py')
arb('الدالة: parse_text_with_stanza(text)')
bullet_arb('Primary: Stanza Pipeline من Stanford NLP (للغة الإنجليزية فقط).')
bullet_arb('المعالجات (Processors): tokenize فقط.')
bullet_arb('Fallback: إذا لم يتوفر Stanza، يُستخدم Regex: split على [.!?\\n].')
bullet_arb('الـ Tokens: تُستخرج عبر regex boundary: \\b\\w[\\w+#.-]*\\b')
bullet_arb('الحد الأقصى للنص المرسل إلى Stanza: 8000 حرف.')

eng('Stanza Output: ParsedText Dataclass')
bullet_eng_arb('sentences', 'قائمة الجمل المستخرجة.')
bullet_eng_arb('tokens', 'قائمة الكلمات المفردة.')
bullet_eng_arb('parser', 'القيمة "stanza" أو "regex_fallback" حسب المتاح.')

eng('2.3 Section Extraction')
arb('الدالة: _extract_sections()')
arb('يتم تحديد أقسام الـ CV بالبحث عن عناوين معروفة:')
bullet_eng_arb('experience', 'Experience, Work History, Employment, خبرات, خبرة')
bullet_eng_arb('education', 'Education, Academic, تعليم, المؤهلات, التعليم')
bullet_eng_arb('projects', 'Projects, Project Experience, مشاريع, المشاريع')
bullet_eng_arb('skills', 'Skills, Technical Skills, مهارات, المهارات التقنية')
bullet_eng_arb('summary', 'Summary, Objective, Profile, About Me, Professional Summary')
bullet_eng_arb('languages', 'Languages, Language Proficiency')
bullet_arb('Fallback: إذا لم يُعثر على أقسام، يتم استخدام Stanza sentences مع keyword matching. مثال: جملة تحتوي على "worked" أو "developed" تُضاف إلى experience.')

eng('2.4 Basic NER (Named Entity Recognition)')
bullet_arb('Name: السطر الأول من النص إذا كان ≤ 5 كلمات و ≤ 60 حرفًا. أو سطر يطابق اسم إنجليزي /^[A-Z][a-z]+ [A-Z][a-z]+/ أو اسم عربي /^[\\u0600-\\u06FF]+ [\\u0600-\\u06FF]+/')
bullet_eng_arb('Email', 'النمط: [\\w.+-]+@[\\w-]+\\.[\\w.-]+')
bullet_eng_arb('Phone', 'النمط: (\\+?\\d[\\d\\s\\-().]{8,}\\d)')
bullet_eng_arb('Location', 'النمط: ^(?:location|address)\\s*[:-]\\s*(.+)$')

eng('2.5 Skill Extraction')
arb('الدالة: extract_catalog_skills() في skill_catalog.py')
bullet_arb('يتم تطبيع النص أولاً.')
bullet_arb('لكل مهارة من قائمة SKILL_KEYWORDS (220+ مهارة) يتم البحث عنها باستخدام regex pattern مخصص.')
bullet_arb('الدالة build_skill_pattern() تُنشئ pattern لكل مهارة مع مراعاة حدود الكلمات.')
bullet_arb('مهارات خاصة لها patterns مخصصة: C++, Vector Database, REST API.')

eng('Input')
arb('raw_text: النص الخام الناتج من المرحلة 1.')

eng('Output')
arb('كائن CandidateProfile الذي يحتوي على:')
bullet_arb('full_name, email, phone, location')
bullet_arb('skills[]: قائمة المهارات المستخرجة')
bullet_arb('experience[], education[], projects[]: نصوص الأقسام')
bullet_arb('languages[]: اللغات')
bullet_arb('total_years_experience: إجمالي سنوات الخبرة (إذا وُجد)')
bullet_arb('raw_text: النص الخام الأصلي')

# ═══════════════════════════════════════════
# STAGE 3: ENHANCED PARSING
# ═══════════════════════════════════════════
doc.add_page_break()
heading_ar('Stage 3 : Enhanced LLM Parsing', 1)

eng('Objective')
arb('استخدام LLM لتحليل أعمق للمهارات: تحديد حالة المهارة (يمتلكها، يتعلمها، لا يمتلكها)، مستوى الخبرة، وسنوات الخبرة. مع التحقق من أن كل استخراج له دليل نصي في الـ CV.')

eng('Entry Point')
arb('الملف: enhanced_cv_parser.py')
arb('الفئة: EnhancedCVParser')
arb('الدالة: parse_async(text)')
arb('الملف المساعد: bilingual_llm.py - الدالة analyze_cv_skills(cv_text)')

eng('3.1 Two-Path Architecture')
arb('يحتوي النظام على مسارين لاستخراج المهارات:')

add_table(
    ['Feature', 'Rule-based', 'LLM-based'],
    [
        ['Speed', 'فوري (بدون طلب خارجي)', 'بطيء (استدعاء LLM)'],
        ['Depth', 'وجود/غياب المهارة فقط', 'Status + Level + Years + Context'],
        ['Dependency', 'لا شيء', 'Ollama / OpenAI'],
        ['Fallback', '-', 'إذا فشل LLM → Rule-based'],
        ['Merge', 'يكون الأساس', 'يُدمج مع Rule-based بأولوية أعلى'],
    ]
)

eng('3.2 Rule-based Extraction')
arb('الدالة: _extract_skills_rule_based()')
bullet_arb('تستخرج المهارات عبر البحث في النص المُطبّع.')
bullet_arb('تستخدم NEGATION_INDICATORS لاكتشاف النفي:')
bullet_arb('"don\'t know", "no experience", "never used", "not familiar", "currently learning"')
bullet_arb('تستخرج سنوات الخبرة عبر pattern: /(\\d+(?:\\.\\d+)?)\\s*(?:year|yr|years|yrs)/')
bullet_arb('تحدد المستوى: < 2 سنوات = Junior, 2-5 = Mid, ≥ 5 = Senior')
bullet_arb('كل مهارة تُرتجع مع: name, level, years, status, confidence, context')

eng('3.3 LLM-based Extraction')
arb('الدالة: _extract_skills_with_llm()')
bullet_arb('تُرسَل السيرة (أول 8000 حرف) إلى LLM عبر prompt مخصص.')
bullet_arb('الـ Prompt: NEGATION_DETECTION_PROMPT في bilingual_llm.py')
bullet_arb('الـ LLM يُعيد JSON بالشكل التالي:')

p = doc.add_paragraph()
r = p.add_run(
    '{\n'
    '  "skills_with_context": [\n'
    '    { "skill": "python", "context": "I have 5 years Python experience",\n'
    '      "status": "has_experience", "years": 5, "level": "senior" }\n'
    '  ],\n'
    '  "negative_skills": ["docker"],\n'
    '  "learning_skills": ["kubernetes"],\n'
    '  "summary": "..."\n'
    '}'
)
r.font.size = Pt(9)
r.font.name = 'Consolas'

eng('3.4 Grounding Verification')
arb('كل مهارة يستخرجها LLM تخضع للتحقق (Grounding).')
bullet_arb('الدالة: _llm_skill_is_grounded() في enhanced_cv_parser.py:522')
bullet_arb('يتم البحث عن اسم المهارة في النص الأصلي للـ CV.')
bullet_arb('إذا وُجد Context، يتم التحقق من وجوده في النص الأصلي.')
bullet_arb('بدون دليل نصي → تُرفض المهارة.')
bullet_arb('هذا يمنع Hallucination من الـ LLM.')

eng('3.5 Merge Strategy')
arb('الدالة: _merge_skill_details()')
bullet_arb('يتم دمج نتائج LLM مع نتائج القواعد.')
bullet_arb('ترتيب الأولوية: HAS_EXPERIENCE > UNKNOWN > LEARNING > NO_EXPERIENCE')
bullet_arb('أي تعارض بين LLM والقواعد → يؤخذ بالأعلى أولوية.')

eng('3.6 Structured Experience Parsing')
arb('الدالة: _parse_experience_entries()')
bullet_arb('يبحث عن أنماط تواريخ وخبرات.')
bullet_arb('يستخرج: title, company, start_date, end_date, description.')
bullet_arb('يحسب total_years_experience تلقائيًا من الفروقات بين التواريخ.')

eng('3.7 Structured Education Parsing')
arb('الدالة: _parse_education_entries()')
bullet_arb('يستخرج: degree, institution, end_date, gpa, description.')
bullet_arb('يكتشف أعلى شهادة (Highest Degree) عبر degree_rank:')
bullet_arb('PhD (7) > Master (5) > Bachelor (4) > Diploma (3) > Associate (2) > Certificate (1)')

eng('Input')
arb('raw_text + CandidateProfile الأساسي من المرحلة 2.')

eng('Output')
arb('CandidateProfile مُعزّز يحتوي بالإضافة إلى:')
bullet_arb('skills_detailed[]: لكل مهارة name, level, years, status, confidence, context, esco_uri, category')
bullet_arb('negative_skills[]: مهارات لا يمتلكها المرشح')
bullet_arb('learning_skills[]: مهارات يتعلمها حالياً')
bullet_arb('experience_entries[]: entries منظمة (title, company, start_date, end_date, description)')
bullet_arb('education_entries[]: entries منظمة (degree, institution, end_date, gpa)')
bullet_arb('highest_degree: أعلى شهادة')
bullet_arb('summary: ملخص نصي من LLM')
bullet_arb('parser_version: "enhanced-v2"')

# ═══════════════════════════════════════════
# STAGE 4: JOB PARSING
# ═══════════════════════════════════════════
doc.add_page_break()
heading_ar('Stage 4 : Job Parsing', 1)

eng('Objective')
arb('تحليل وصف الوظيفة لاستخراج المهارات المطلوبة، المهارات الاختيارية، ومستوى الأقدمية.')

eng('Entry Point')
arb('الملف: job_parser.py')
arb('الدالة: parse_job_description(text)')

eng('4.1 Skill Extraction by Section')
bullet_eng_arb('Required Skills', 'يُبحث عن أقسام: Requirements, Must Have, Required Skills, Required Qualifications.')
bullet_eng_arb('Optional Skills', 'يُبحث عن أقسام: Nice to Have, Preferred, Bonus, Desirable.')
bullet_eng_arb('Bonus Detection', 'نص يحتوي على "is a plus" أو "nice to have" يُضاف إلى optional.')

eng('4.2 Fallback Strategy')
bullet_arb('إذا لم يُعثر على أقسام، تُستخرج كل المهارات من النص بأكمله.')
bullet_arb('يتم إزالة المهارات المكررة والمهارات غير المقبولة.')

eng('4.3 Seniority Detection')
arb('الدالة: _detect_seniority()')
bullet_arb('Junior: الكلمات junior, entry level, associate')
bullet_arb('Mid: الكلمات mid, intermediate')
bullet_arb('Senior: الكلمات senior, lead, principal, staff')

eng('4.4 Title Extraction')
bullet_arb('أول سطر غير فارغ في النص (أقصى 120 حرفًا).')

eng('Input')
arb('وصف الوظيفة كنص (string).')

eng('Output')
arb('كائن JobProfile يحتوي على:')
bullet_arb('title: عنوان الوظيفة')
bullet_arb('description: النص الأصلي للوصف')
bullet_arb('required_skills[]: المهارات المطلوبة (مرتبة، بدون تكرار)')
bullet_arb('optional_skills[]: المهارات الاختيارية (مستبعد منها المهارات المطلوبة)')
bullet_arb('seniority: واحد من junior, mid, senior, lead, principal, staff أو None')

# ═══════════════════════════════════════════
# STAGE 5: MATCHING (THE CORE)
# ═══════════════════════════════════════════
doc.add_page_break()
heading_ar('Stage 5 : Matching Engine', 1)

eng('Objective')
arb('حساب درجة تطابق عددية (Score) بين 0.0 و 1.0 لكل مرشح مع الوظيفة. تعتمد الدرجة على 5 مكونات رئيسية.')

eng('Entry Point')
arb('الملف الرئيسي: hybrid_matcher.py')
arb('الفئة: HybridMatchingEngine')
arb('الدالة: match()')
arb('الملفات المساعدة: matching.py, embedding.py, vector_store.py, ollama_cross_encoder.py')

eng('5.1 Step 1: Job Embedding')
bullet_arb('نص الوظيفة = job.title + " " + job.description.')
bullet_arb('يتم إرسال النص إلى Embedding Service للحصول على متجه (Vector).')
bullet_arb('إذا كان embedding_provider = "hash" → يُلغى التشابه الدلالي (Semantic Score = 0).')
bullet_arb('إذا كان النص قصيرًا جدًا → تعطيل Semantic Score.')

eng('5.2 Step 2: Candidate Embeddings (Batch)')
bullet_arb('يُبنى نص تمثيلي لكل مرشح: Skills: ..., Experience: ..., Education: ..., Projects: ..., CV text: ...')
bullet_arb('الدالة: build_candidate_embedding_text() في candidate_text.py.')
bullet_arb('البحث في Vector Store عن Embeddings مخزنة مسبقًا.')
bullet_arb('للمرشحين الجدد → حساب Embeddings دفعة واحدة (Batch).')
bullet_arb('تخزين النتائج في Vector Store للاستخدام المستقبلي.')

eng('5.3 Step 3: Semantic Similarity')
bullet_arb('تحويل متجهات المرشحين إلى مصفوفة NumPy.')
bullet_arb('حساب Cosine Similarity دفعة واحدة: similarity = dot(C, J) / (||C|| * ||J||)')
bullet_arb('Clipping إلى [0, 1].')

eng('5.4 Step 4: Skill Matching')
arb('الدالة: _compute_skill_match()')
arb('لكل مهارة مطلوبة، يتم البحث بالترتيب التالي:')

add_table(
    ['Priority', 'Match Type', 'Description', 'Confidence'],
    [
        ['1', 'Exact Match', 'اسم المهارة نفسه موجود في Skills المرشح (بعد التطبيع)', '1.00'],
        ['2', 'Synonym Match', 'عبر SYNONYM_MAP (Docker = Containerization)', '0.85'],
        ['3', 'ESCO Match', 'عبر URI موحد في ESCO Taxonomy', '0.95'],
        ['4', 'ESCO Related', 'عبر علاقات ESCO: broader / narrower / related', '0.60 - 0.95'],
        ['5', 'Text Evidence', 'البحث عن المهارة في النص الخام للـ CV', '0.80'],
    ]
)

eng('Skill Negation Detection')
bullet_arb('إذا وجدت عبارات نفي في سياق المهارة → لا تُحتسب.')
bullet_arb('قائمة NEGATION_INDICATORS: "don\'t know", "no experience", "not familiar", "never used", "currently learning".')
bullet_arb('الدالة: _candidate_skill_is_negated() في hybrid_matcher.py:975')

eng('Scoring Formulas')
arb('Required Score = sum(confidence لكل matched_required) / total_required')
arb('Optional Score = sum(confidence لكل matched_optional) / total_optional')
arb('Skill Score = 0.8 * Required Score + 0.2 * Optional Score')

eng('5.5 Seniority & Experience Scoring')
arb('Years Score = min(1.0, estimated_years / 10)')

arb('Seniority Score:')
arb('Seniority mapping:')
bullet_arb('Junior: (0-2 years)')
bullet_arb('Mid: (2-5 years)')
bullet_arb('Senior: (5-10 years)')
bullet_arb('Lead: (8-15 years)')
bullet_arb('Principal / Staff: (10-20 years)')
bullet_arb('إذا كانت سنوات المرشح ضمن النطاق → Score = 1.0')
bullet_arb('إذا كانت أقل → Score = max(0.0, years / min_years * 0.5)')
bullet_arb('إذا كانت أكثر → Penalty = min(0.3, excess * 0.05) → Score = max(0.5, 1.0 - penalty)')

eng('5.6 Final Score Formula')
arb('Base Weights (توزانها يساوي 1.0):')
arb('Final Score = (0.55 x Required Skills) + (0.20 x Optional Skills) + (0.15 x Semantic) + (0.05 x Experience) + (0.05 x Seniority)')

arb('Score Cap (سقف للدرجة):')
bullet_arb('Required Score = 0% → Score Cap = 0.40')
bullet_arb('Required Score < 50% → Score Cap = 0.55')
bullet_arb('Required Score < 75% → Score Cap = 0.75')
bullet_arb('Required Score < 90% → Score Cap = 0.85')
bullet_arb('Required Score < 100% → Score Cap = 0.90')
bullet_arb('Required Score = 100% → Score Cap = 1.00 (بدون سقف)')
bullet_arb('Final = min(Pre-Cap Score, Score Cap)')

eng('5.7 Cross-Encoder Re-ranking (Optional)')
arb('الدالة: _cross_encoder_rerank()')
bullet_arb('يُمرر أفضل N مرشح إلى OllamaCrossEncoder.')
bullet_arb('OllamaCrossEncoder يستخدم LLM لتقييم زوج (Job, Candidate).')
bullet_arb('يرجع Score بين 0.0 و 1.0.')
bullet_arb('التعديل محدود بـ ±0.05 كحد أقصى.')
bullet_arb('Cross Encoder Weight = 0.25.')
bullet_arb('Formula: adjustment = clamp(0.25 x (cross_score - base_score), -0.05, +0.05).')

eng('Input')
arb('JobProfile (من Stage 4) + CandidateProfile[] (من Stage 3)')

eng('Output: HybridMatchResult')
bullet_arb('candidate_id, final_score')
bullet_arb('skill_match: matched_required[], matched_optional[], missing_required[], required_score, optional_score')
bullet_arb('semantic_score, cross_encoder_score (أو None)')
bullet_arb('reasoning: scoring_model, scoring_formula, score_breakdown, strengths[], gaps[], recommendations[]')

# ═══════════════════════════════════════════
# STAGE 6: REPORT
# ═══════════════════════════════════════════
doc.add_page_break()
heading_ar('Stage 6 : Explainability Report', 1)

eng('Objective')
arb('ترجمة نتائج الـ Matching إلى تقرير بشري مفهوم مع تحليل الفجوات والتوصيات.')

eng('Entry Point')
arb('الملف: explainability.py')
arb('الدالة: generate_candidate_report(session, job_id, candidate_id)')

eng('6.1 Score Breakdown')
bullet_eng_arb('similarity_score', 'التشابه الدلالي بين الـ Job والـ Candidate.')
bullet_eng_arb('required_skills_score', 'نسبة المهارات المطلوبة المُحققة.')
bullet_eng_arb('optional_skills_score', 'نسبة المهارات الاختيارية المُحققة.')
bullet_eng_arb('overall_score', 'الدرجة النهائية (من MatchResult أو يُعاد حسابها).')
bullet_eng_arb('pre_cap_score', 'الدرجة قبل تطبيق السقف.')
bullet_eng_arb('score_cap', 'السقف المطبق.')
bullet_eng_arb('score_cap_reason', 'شرح نصي عن سبب السقف.')

eng('6.2 Skill Gap Analysis')
bullet_eng_arb('matched_required[]', 'المهارات المطلوبة الموجودة عند المرشح.')
bullet_eng_arb('missing_required[]', 'المهارات المطلوبة المفقودة (أهم جزء في التقرير).')
bullet_eng_arb('matched_optional[]', 'المهارات الاختيارية الموجودة.')
bullet_eng_arb('items[]', 'قائمة بكل المهارات مع خاصية required/matched.')

eng('6.3 Strengths and Weaknesses')
bullet_arb('Strengths: المهارات المطلوبة المُحققة (أول 5).')
bullet_arb('Weaknesses: المهارات المطلوبة المفقودة.')

eng('6.4 Recommendation (Text)')
bullet_arb('Score ≥ 0.80: "Highly recommended. Excellent match for the position."')
bullet_arb('Score ≥ 0.60: "Recommended. Good overall fit with some areas for development."')
bullet_arb('Score ≥ 0.40: "Consider with reservations. Meets basic requirements but has gaps."')
bullet_arb('Score < 0.40: "Not recommended. Significant gaps in required qualifications."')

eng('6.5 Interview Score Blend')
bullet_arb('إذا وُجدت مقابلة مكتملة → Blended Score = (0.65 x CV_match) + (0.35 x Interview_score).')
bullet_arb('المنطق في interview_analysis.py:136.')
bullet_arb('Formula: "0.65 CV/job match + 0.35 post-interview answer analysis".')

eng('Input')
arb('job_id + candidate_id (للبحث في قاعدة البيانات عن Job, Candidate, MatchResult).')

eng('Output: CandidateReportResponse')
bullet_arb('job_title, candidate_name')
bullet_arb('score_breakdown: كل التفاصيل الرقمية')
bullet_arb('skill_gap: المهارات المطابقة والمفقودة')
bullet_arb('strengths[], weaknesses[]')
bullet_arb('recommendation: نص التوصية')

# ═══════════════════════════════════════════
# EXAMPLE
# ═══════════════════════════════════════════
doc.add_page_break()
heading_ar('Example : Jad + AI Engineer', 1)

arb('تطبيق الـ Pipeline الكامل على مثال حقيقي.')

eng('Candidate: Jad Ahmad')
bullet_arb('5 سنوات خبرة.')
bullet_arb('المهارات: Python, PyTorch, FastAPI, NLP, Machine Learning, PostgreSQL.')
bullet_arb('مهارات يتعلمها: Docker, Kubernetes.')
bullet_arb('خبرة: Senior Backend Developer (2020-2025), Junior Developer (2018-2020).')

eng('Job: AI Engineer')
bullet_arb('Required Skills: Python, Machine Learning, NLP, PyTorch.')
bullet_arb('Optional Skills: Docker, Kubernetes, FastAPI.')
bullet_arb('Seniority: Mid.')

eng('Stage 5 - Matching Calculation')
add_table(
    ['Component', 'Details', 'Score'],
    [
        ['Required Skills', 'Python / ML / NLP / PyTorch → 4/4', '1.000'],
        ['Optional Skills', 'FastAPI only → 1/3', '0.333'],
        ['Semantic Similarity', 'Cosine(Job Vector, Jad Vector)', '~0.82'],
        ['Experience', '5 years / 10', '0.500'],
        ['Seniority', 'Mid (2-5) → 5 years → exact match', '1.000'],
    ]
)

arb('Final Score Formula:')
arb('Pre-Cap = (0.55 x 1.0) + (0.20 x 0.333) + (0.15 x 0.82) + (0.05 x 0.5) + (0.05 x 1.0) = 0.815')
arb('Required Score = 1.0 → Score Cap = 1.0 (no cap)')
arb('Final Score = min(0.815, 1.0) = 0.815')

eng('Stage 6 - Report')
bullet_arb('Strengths: Python, Machine Learning, NLP, PyTorch, FastAPI.')
bullet_arb('Weaknesses: Docker, Kubernetes (Optional skills missing but in learning).')
bullet_arb('Recommendation: "Recommended. Good overall fit with some areas for development."')

# ═══════════════════════════════════════════
# FILES REFERENCE
# ═══════════════════════════════════════════
doc.add_page_break()
heading_ar('Files Reference', 1)
arb('جميع الملفات المتعلقة بالـ NLP في المشروع:')

add_table(
    ['File Path', 'Main Classes / Functions', 'NLP Role'],
    [
        ['services/cv_parser.py', 'extract_text(), parse_cv_text()', 'Text extraction + basic parsing'],
        ['services/stanza_nlp.py', 'parse_text_with_stanza(), split_sentences_with_stanza()', 'Sentence splitting, Tokenization'],
        ['services/enhanced_cv_parser.py', 'EnhancedCVParser, parse_async()', 'Enhanced LLM-based parsing'],
        ['services/skill_catalog.py', 'SKILL_KEYWORDS, SYNONYM_MAP, build_skill_pattern()', 'Skill ontology + pattern matching'],
        ['services/embedding.py', 'LocalEmbeddingService, OllamaEmbeddingService, HashEmbeddingService', 'Embedding generation'],
        ['services/vector_store.py', 'VectorStore, query_similar(), upsert_embedding()', 'Vector storage + similarity search'],
        ['services/hybrid_matcher.py', 'HybridMatchingEngine, match()', 'Core matching algorithm'],
        ['services/matching.py', 'rank_candidates(), compute_skill_score()', 'Legacy matching interface'],
        ['services/job_parser.py', 'parse_job_description()', 'Job description analysis'],
        ['services/bilingual_llm.py', 'BilingualLLMService, evaluate_answer(), analyze_cv_skills()', 'LLM integration'],
        ['services/ollama_cross_encoder.py', 'OllamaCrossEncoder, predict()', 'Cross-encoder re-ranking'],
        ['services/esco_service.py', 'ESCOSkillService, normalize_skill(), get_related_skills()', 'ESCO taxonomy integration'],
        ['services/esco_extractor.py', 'EscoSkillExtractor, extract_skills()', 'ESCO API + embedding'],
        ['services/explainability.py', 'generate_candidate_report(), compare_candidates()', 'Report generation'],
        ['services/interview.py', 'build_grounded_question_items(), _evaluate_single_answer()', 'Interview engine'],
        ['services/enhanced_interview.py', 'EnhancedInterviewService, evaluate_answer_with_llm()', 'LLM interview evaluation'],
        ['services/interview_analysis.py', 'analyze_completed_interview(), upsert_interview_match_result()', 'Post-interview analysis'],
        ['services/voice_service.py', 'VoiceService, process_audio()', 'Speech-to-Text, Text-to-Speech'],
        ['services/rag.py', 'ingest_knowledge_base(), query_knowledge()', 'RAG retrieval'],
        ['services/candidate_text.py', 'build_candidate_embedding_text()', 'Candidate text builder for embeddings'],
        ['services/project_semantic.py', 'compute_junior_project_semantic_bonus()', 'Junior project bonus'],
        ['schemas/candidate.py', 'CandidateProfile, SkillDetail, SkillStatus, SkillLevel', 'Data schemas'],
        ['schemas/job.py', 'JobProfile', 'Job data schema'],
        ['schemas/report.py', 'CandidateReportResponse, ScoreBreakdown, SkillGapAnalysis', 'Report schemas'],
    ]
)

doc.add_paragraph()
doc.add_paragraph()
eng('- - -', size=14, bold=False)
eng('End of Document', size=14)

# ── Save ──
output_path = os.path.join(os.path.dirname(__file__), 'NLP_Pipeline_Clean.docx')
doc.save(output_path)
print(f'Document saved to: {output_path}')
